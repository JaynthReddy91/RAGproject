import os
import re
from pathlib import Path
from typing import List, Dict, Any, Optional

class Document:
    def __init__(self, text: str, metadata: Dict[str, Any]):
        self.text = text
        self.metadata = metadata

    def __repr__(self):
        return f"Document(source={self.metadata.get('source')}, page={self.metadata.get('page')}, len={len(self.text)})"

class DocumentLoader:
    def __init__(self):
        pass

    def clean_text(self, text: str) -> str:
        """Cleans extracted text by standardizing spaces, line breaks, and simple header/footer patterns."""
        if not text:
            return ""
        
        # Replace multiple spaces with a single space
        text = re.sub(r'[ \t]+', ' ', text)
        
        # Replace 3 or more newlines with double newline
        text = re.sub(r'\n{3,}', '\n\n', text)
        
        # Remove trailing/leading spaces on lines
        text = '\n'.join([line.strip() for line in text.split('\n')])
        
        # Basic removal of common header/footer relics (like lonely numbers at the beginning or end of page text)
        lines = text.split('\n')
        if lines:
            # Check if first or last line is just a number (page numbering)
            if re.match(r'^\d+$', lines[0].strip()):
                lines = lines[1:]
            if lines and re.match(r'^\d+$', lines[-1].strip()):
                lines = lines[:-1]
        
        return '\n'.join(lines).strip()

    def load_pdf(self, file_path: Path) -> List[Document]:
        """Loads PDF using PyMuPDF (fitz), extracting page number and text metadata."""
        import fitz  # PyMuPDF
        
        documents = []
        try:
            doc = fitz.open(file_path)
            total_pages = len(doc)
            
            for page_idx in range(total_pages):
                page = doc.load_page(page_idx)
                page_num = page_idx + 1
                text = page.get_text()
                
                # Check for scanned pages / lack of text
                if not text.strip():
                    # Check if there are images, indicating it might be a scanned page
                    images = page.get_images()
                    if images:
                        text = f"[Scanned page image metadata: {len(images)} images detected. Text extraction empty.]"
                    else:
                        text = "[Empty page]"

                cleaned_text = self.clean_text(text)
                
                metadata = {
                    "source": file_path.name,
                    "file_path": str(file_path),
                    "page": page_num,
                    "total_pages": total_pages,
                    "file_type": "pdf"
                }
                documents.append(Document(cleaned_text, metadata))
                
        except Exception as e:
            print(f"Error loading PDF {file_path}: {e}")
            
        return documents

    def load_docx(self, file_path: Path) -> List[Document]:
        """Loads Microsoft Word files using python-docx."""
        import docx
        
        documents = []
        try:
            doc = docx.Document(file_path)
            
            # Since DOCX doesn't have native "pages" like PDF, we group text paragraphs
            # and tables into logical "sections" or virtual pages (e.g. every 500 words or 5 paragraphs)
            content_blocks = []
            
            # Extract paragraphs
            for para in doc.paragraphs:
                p_text = para.text.strip()
                if p_text:
                    content_blocks.append(p_text)
                    
            # Extract tables and format rows as Markdown-style cell rows
            for table in doc.tables:
                table_text = []
                for row in table.rows:
                    row_cells = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
                    table_text.append(" | ".join(row_cells))
                if table_text:
                    content_blocks.append("\n".join(table_text))
            
            # Partition blocks into virtual pages (approx 5 paragraphs or table blocks per page)
            chunk_size = 5
            virtual_page = 1
            for i in range(0, len(content_blocks), chunk_size):
                sub_blocks = content_blocks[i:i + chunk_size]
                text_content = "\n\n".join(sub_blocks)
                cleaned_text = self.clean_text(text_content)
                
                metadata = {
                    "source": file_path.name,
                    "file_path": str(file_path),
                    "page": virtual_page,
                    "total_pages": (len(content_blocks) + chunk_size - 1) // chunk_size,
                    "file_type": "docx"
                }
                documents.append(Document(cleaned_text, metadata))
                virtual_page += 1
                
        except Exception as e:
            print(f"Error loading DOCX {file_path}: {e}")
            
        return documents

    def load_txt(self, file_path: Path) -> List[Document]:
        """Loads plain text files."""
        documents = []
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                text = f.read()
            
            cleaned_text = self.clean_text(text)
            
            # Similar to DOCX, split into virtual pages if file is very long (e.g., every 2500 characters)
            chars_per_page = 2500
            total_len = len(cleaned_text)
            total_pages = (total_len + chars_per_page - 1) // chars_per_page
            
            if total_pages <= 1:
                metadata = {
                    "source": file_path.name,
                    "file_path": str(file_path),
                    "page": 1,
                    "total_pages": 1,
                    "file_type": "txt"
                }
                documents.append(Document(cleaned_text, metadata))
            else:
                for idx in range(total_pages):
                    start = idx * chars_per_page
                    end = min(start + chars_per_page, total_len)
                    page_text = cleaned_text[start:end]
                    metadata = {
                        "source": file_path.name,
                        "file_path": str(file_path),
                        "page": idx + 1,
                        "total_pages": total_pages,
                        "file_type": "txt"
                    }
                    documents.append(Document(page_text, metadata))
                    
        except Exception as e:
            print(f"Error loading TXT {file_path}: {e}")
            
        return documents

    def load_file(self, file_path: Path) -> List[Document]:
        """Loads any supported file type based on its file extension."""
        ext = file_path.suffix.lower()
        if ext == ".pdf":
            return self.load_pdf(file_path)
        elif ext == ".docx":
            return self.load_docx(file_path)
        elif ext in [".txt", ".md"]:
            return self.load_txt(file_path)
        else:
            print(f"Skipping unsupported file type: {file_path}")
            return []

    def load_directory(self, dir_path: Path) -> List[Document]:
        """Loads all supported files from the directory."""
        all_documents = []
        if not dir_path.exists() or not dir_path.is_dir():
            print(f"Directory does not exist: {dir_path}")
            return all_documents
        
        for file in dir_path.iterdir():
            if file.is_file():
                all_documents.extend(self.load_file(file))
                
        return all_documents
